import os
from docx import Document
from docx.shared import Inches
#from docx2pdf import convert

def collect_and_insert_images(category_path, plot_name, output_folder):
    s_age =''

    # Iterate through subject folders within the category
    for subject_folder in os.listdir(category_path):
        print(f'Processing {subject_folder}')
        subject_path = os.path.join(category_path, subject_folder)
        if os.path.isdir(subject_path):
            if subject_folder != '.ipynb_checkpoints':

              extrat_info= subject_folder.split('_')
              subject_id = extrat_info[0]
              s_age = extrat_info[1]

              doc_path = f'{subject_id}_{s_age}.docx'
              p_name = plot_name.split('_')[1]
              # Save the Word document with the subject ID as part of the name
              if '_R' in subject_folder:
                  doc_path = os.path.join(output_folder, f'{subject_id}_{s_age}_R.docx')
                  heading = f'{p_name}_{subject_id}_{s_age}_R'
                
              else:
                  doc_path = os.path.join(output_folder, f'{subject_id}_{s_age}.docx')
                  heading = f'{p_name}_{subject_id}_{s_age}'

              if os.path.exists(doc_path):
                  doc = Document(doc_path)
              else:
                  # Create a new Word document for each subject
                  doc = Document()
              
             
              # Add a heading with the subject name
              doc.add_heading(heading, level=1)

              # Sort PNG files based on 'top' or 'bot' in the filename
              imu_files = sorted(os.listdir(subject_path), key=lambda x: ('bot' in x, x))

              # Iterate through sorted PNG files in the subject folder
              for file_name in imu_files:
                  if file_name.lower().endswith(".png") and file_name.startswith(plot_name):
                      # Add image to the document
                      image_path = os.path.join(subject_path, file_name)
                      doc.add_picture(image_path, width=Inches(7))  # Adjust width as needed
              doc.save(doc_path)

              # Convert the Word document to PDF
              #pdf_path = os.path.join(pdf_output_folder, f'{subject_id}_{s_age}.pdf')
              #convert(doc_path, pdf_path)

        # Create a new section for each subject
        doc.add_section()


plots_name = ['plot_FL0','plot_FL1','plot_FL2', 'plot_FL3', 'plot_FL4', 'plot_CH1', 'plot_CH2', 'plot_CR1', 'plot_CR3a','plot_CR3b','plot_CR4a','plot_CR4b']
data_folder_path = '../plots/'

rated Word documents
output_folder_path = '../plots/subject_based/'
os.makedirs(output_folder_path, exist_ok=True)


categories =['EBI/', 'EMG/','IMU/']

for plot in plots_name:
  print('Plotting: ', plot)
  for cat in categories:
    base_folder= data_folder_path + cat
    collect_and_insert_images(base_folder, plot,output_folder_path)
